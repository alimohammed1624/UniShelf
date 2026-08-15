from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from bs4 import BeautifulSoup
import re
import sys

def add_code_block(doc, text):
    """Add a code block with monospace font and dark background"""
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.5)
    run = p.add_run(text.strip())
    run.font.name = 'Courier New'
    run.font.size = Pt(9)
    
    # Add shading for code block background
    from docx.oxml.ns import qn
    shading_elm = p._element.get_or_add_pPr()
    shd = shading_elm.makeelement(qn('w:shd'), {
        qn('w:fill'): '1A1A2E',
        qn('w:val'): 'clear'
    })
    shading_elm.append(shd)

def convert_table_to_docx(table_elem, doc):
    """Convert an HTML table to DOCX table"""
    headers = table_elem.find_all('th')
    rows = table_elem.find_all('tr')
    
    if not rows:
        return
    
    # Determine number of columns
    num_cols = len(headers) if headers else 0
    if num_cols == 0:
        first_row = rows[0]
        num_cols = len(first_row.find_all(['td', 'th']))
    
    if num_cols == 0:
        return
    
    doc_table = doc.add_table(rows=1, cols=num_cols)
    doc_table.style = 'Table Grid'
    
    # Add header row if exists
    if headers:
        header_row = doc_table.rows[0]
        for i, th in enumerate(headers):
            if i < num_cols:
                cell = header_row.cells[i]
                cell.text = th.get_text().strip()
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.bold = True
    
    # Add data rows from tbody or remaining rows
    tbody = table_elem.find('tbody')
    if tbody:
        data_rows = tbody.find_all('tr', recursive=False)
    else:
        data_rows = [r for r in rows[1:] if r.find_parent('thead') is None]
    
    for row_elem in data_rows:
        row = doc_table.add_row()
        cells = row_elem.find_all(['td', 'th'])
        
        for i, cell_elem in enumerate(cells):
            if i < num_cols:
                cell = row.cells[i]
                code_elem = cell_elem.find('code')
                if code_elem:
                    text = code_elem.get_text()
                    p = cell.paragraphs[0]
                    run = p.add_run(text)
                    run.font.name = 'Courier New'
                    run.font.size = Pt(9)
                else:
                    cell.text = cell_elem.get_text().strip()

def process_element(element, doc):
    """Recursively process HTML elements and add to DOCX"""
    if isinstance(element, str) or not hasattr(element, 'name'):
        return
    
    tag = element.name
    class_name = element.get('class', [])
    
    # Skip non-content tags
    if tag in ['html', 'head', 'body', 'div', 'span']:
        for child in element.children:
            process_element(child, doc)
        return
    
    if tag == 'h1':
        p = doc.add_heading(level=0)
        run = p.add_run(element.get_text().strip())
        
    elif tag == 'h2':
        p = doc.add_heading(level=1)
        run = p.add_run(element.get_text().strip())
        
    elif tag == 'h3':
        p = doc.add_heading(level=2)
        run = p.add_run(element.get_text().strip())
        
    elif tag == 'p':
        text = element.get_text().strip()
        if text:
            # Check for subtitle class
            if 'subtitle' in class_name:
                p = doc.add_paragraph()
                run = p.add_run(text)
                run.font.color.rgb = RGBColor(0xA0, 0xA0, 0xB0)
                run.font.size = Pt(14)
                doc.add_paragraph()  # spacer
            else:
                doc.add_paragraph(text)
                
    elif tag == 'pre':
        code_block = element.find('code')
        code_text = code_block.get_text() if code_block else element.get_text()
        add_code_block(doc, code_text)
        
    elif tag == 'table':
        convert_table_to_docx(element, doc)
        
    elif tag in ['ul', 'ol']:
        for li in element.find_all('li', recursive=False):
            text = li.get_text().strip()
            if text:
                doc.add_paragraph(text, style='List Bullet')

def convert_html_to_docx(html_file_path, docx_file_path):
    """Convert an HTML file to DOCX format"""
    
    with open(html_file_path, 'r', encoding='utf-8') as f:
        html_content = f.read()
    
    # Parse HTML
    soup = BeautifulSoup(html_content, 'html.parser')
    
    # Create document
    doc = Document()
    
    # Set default font
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)
    
    # Process body content
    body = soup.find('body')
    if not body:
        print("No body found in HTML")
        return
    
    for element in body.children:
        process_element(element, doc)
    
    # Save document
    doc.save(docx_file_path)

if __name__ == '__main__':
    if len(sys.argv) != 3:
        print("Usage: python convert_html_to_docx.py <html_file> <docx_file>")
        sys.exit(1)
    
    html_file = sys.argv[1]
    docx_file = sys.argv[2]
    
    try:
        convert_html_to_docx(html_file, docx_file)
        print(f"Successfully converted {html_file} to {docx_file}")
    except Exception as e:
        print(f"Error converting file: {e}")
        import traceback
        traceback.print_exc()
